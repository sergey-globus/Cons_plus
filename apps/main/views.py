from django.shortcuts import render, redirect
from django.contrib import messages
from django.http import HttpResponse, Http404
from .models import FAQ, Consultation, DocumentTemplate
from .forms import ConsultationForm, DocumentGeneratorForm
from django.conf import settings
from io import BytesIO
import os
from docx import Document
from docxtpl import DocxTemplate
import mammoth
from django.db.models import Q
from django.shortcuts import get_object_or_404
from django.core.mail import send_mail
from django.urls import reverse
from django.core.files import File
import tempfile
import asyncio
from asgiref.sync import sync_to_async
from concurrent.futures import ThreadPoolExecutor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Global thread pool for synchronous IO operations
executor = ThreadPoolExecutor(max_workers=10)

# Helper functions for session operations
def session_set(session, key, value):
    session[key] = value
    session.save()

def session_get(session, key, default=None):
    return session.get(key, default)

async def index(request):
    return await sync_to_async(render)(request, 'main/index.html')

async def consultation(request):
    if request.method == 'POST':
        form = ConsultationForm(request.POST)
        if form.is_valid():
            try:
                # Save consultation in a separate task
                save_task = asyncio.create_task(sync_to_async(form.save)())
                consultation = await save_task
                
                # Prepare email data
                first_name = consultation.first_name
                last_name = consultation.last_name
                email = consultation.email
                
                subject = 'Ваша заявка на консультацию принята'
                message = f'''
                Уважаемый(ая) {first_name} {last_name},
            
                Благодарим вас за обращение на наш сайт - МаркетНадзор!
                Ваш вопрос: "{consultation.question[:50]}..."
            
                Наши юристы рассмотрят вашу заявку в ближайшее время и свяжутся с вами в течение 24 часов, в том числе в выходные и праздничные дни.
            
                С уважением,
                Команда МаркетНадзор
                '''
                
                # Send email in a separate task
                email_task = asyncio.create_task(
                    sync_to_async(send_mail)(
                        subject=subject,
                        message=message,
                        from_email=settings.DEFAULT_FROM_EMAIL,
                        recipient_list=[email],
                        fail_silently=False,
                    )
                )
                await email_task
                
                messages.success(request, 'Ваша заявка принята! На ваш email отправлено подтверждение.')
                request.session['show_success_gif'] = True
                
                return redirect('consultation')
            
            except Exception as e:
                return redirect('consultation')
    else:
        form = ConsultationForm()
    
    return await sync_to_async(render)(request, 'main/consultation.html', {'form': form})

async def document_result(request, template_id):
    # Get session data in a separate task
    session_task = asyncio.create_task(sync_to_async(session_get)(request.session, 'form_data', {}))
    form_data = await session_task
    
    # Get template in a separate task
    template_task = asyncio.create_task(sync_to_async(DocumentTemplate.objects.get)(id=template_id))
    template = await template_task

    document_html = "<p>Файл шаблона не загружен.</p>"

    if template.word_document:
        def process_template():
            tmp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp_path = tmp.name
                    doc = DocxTemplate(template.word_document.path)
                    doc.render(form_data if form_data else {})
                    doc.save(tmp_path)

                with open(tmp_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    return result.value
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        # Process template in thread pool
        process_task = asyncio.get_event_loop().run_in_executor(executor, process_template)
        document_html = await process_task

    return await sync_to_async(render)(
        request, 
        "main/document_result.html", 
        {
            "template_name": template.name,
            "document_html": document_html,
            "template_id": template.id,
        }
    )

async def preview_template(request, template_id):
    template_task = asyncio.create_task(sync_to_async(get_object_or_404)(DocumentTemplate, id=template_id))
    template = await template_task
    
    document_html = ""

    if template.word_template:
        def convert_to_html():
            with open(template.word_template.path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                return result.value
        
        convert_task = asyncio.get_event_loop().run_in_executor(executor, convert_to_html)
        document_html = await convert_task

    context = {
        'template': template,
        'document_html': document_html,
        'translated_fields': template.get_translated_fields(),
    }
    return await sync_to_async(render)(request, "main/preview_template.html", context)

async def document_generator(request):
    # Get templates in a separate task
    templates_task = asyncio.create_task(sync_to_async(list)(DocumentTemplate.objects.all()))
    templates = await templates_task

    # Handle template selection
    selected_template_id = request.GET.get('template_id')
    try:
        selected_template_id_int = int(selected_template_id)
    except (TypeError, ValueError):
        selected_template_id_int = None

    if request.method == 'POST':
        template_id = request.POST.get('template_id')
        if template_id:
            try:
                # Get template in a separate task
                template_task = asyncio.create_task(sync_to_async(DocumentTemplate.objects.get)(id=template_id))
                template = await template_task

                # Prepare form data
                form_data = {}
                for k, v_list in request.POST.lists():
                    if k in ('csrfmiddlewaretoken', 'template_id'):
                        continue
                    value = next((v for v in v_list if v.strip()), '')
                    form_data[k] = value if value else f'[{k}]'

                # Save to session in a separate task
                session_task = asyncio.create_task(sync_to_async(session_set)(request.session, 'form_data', form_data))
                await session_task

                return redirect(reverse('document_result', kwargs={'template_id': template.id}))

            except DocumentTemplate.DoesNotExist:
                await sync_to_async(messages.error)(request, 'Выбранный шаблон не найден')

    return await sync_to_async(render)(request, 'main/document_generator.html', {
        'templates': templates,
        'selected_template_id': selected_template_id_int,
    })

async def download_pdf(request):
    document_text = request.session.get('generated_document', '')
    template_name = request.session.get('template_name', 'Документ')

    if not document_text:
        messages.error(request, 'Документ не найден. Создайте документ заново.')
        return redirect('document_generator')

    def generate_pdf():
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        try:
            font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                font_name = 'DejaVuSans'
            else:
                font_name = 'Helvetica'
        except:
            font_name = 'Helvetica'

        p.setFont(font_name, 16)
        p.drawString(50, height - 50,
                    template_name.encode('utf-8').decode('utf-8') if font_name == 'Helvetica' else template_name)

        p.setFont(font_name, 12)
        lines = document_text.split('\n')
        y_position = height - 100

        for line in lines:
            if y_position < 50:
                p.showPage()
                p.setFont(font_name, 12)
                y_position = height - 50

            if len(line) > 80:
                words = line.split(' ')
                current_line = ''
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + ' '
                    else:
                        if current_line:
                            try:
                                p.drawString(50, y_position, current_line.encode('utf-8').decode(
                                    'utf-8') if font_name == 'Helvetica' else current_line)
                            except:
                                p.drawString(50, y_position, current_line)
                            y_position -= 15
                        current_line = word + ' '
                if current_line:
                    try:
                        p.drawString(50, y_position, current_line.encode('utf-8').decode(
                            'utf-8') if font_name == 'Helvetica' else current_line)
                    except:
                        p.drawString(50, y_position, current_line)
                    y_position -= 15
            else:
                try:
                    p.drawString(50, y_position, line.encode('utf-8').decode('utf-8') if font_name == 'Helvetica' else line)
                except:
                    p.drawString(50, y_position, line)
                y_position -= 15

        p.save()
        buffer.seek(0)
        return buffer

    # Generate PDF in thread pool
    generate_task = asyncio.get_event_loop().run_in_executor(executor, generate_pdf)
    buffer = await generate_task

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{template_name}.pdf"'
    return response

async def download_word(request, template_id):
    # Get template and session data in parallel
    template_task = asyncio.create_task(sync_to_async(get_object_or_404)(DocumentTemplate, id=template_id))
    session_task = asyncio.create_task(sync_to_async(session_get)(request.session, 'form_data', {}))
    
    template, form_data = await asyncio.gather(template_task, session_task)

    if not template.word_document:
        await sync_to_async(messages.error)(request, 'Шаблон документа не найден.')
        return redirect('document_generator')

    def generate_docx():
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp_path = tmp.name
                doc = DocxTemplate(template.word_document.path)
                doc.render(form_data)
                doc.save(tmp_path)

            with open(tmp_path, 'rb') as f:
                return f.read()
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)

    try:
        # Generate document in thread pool
        content_task = asyncio.get_event_loop().run_in_executor(executor, generate_docx)
        content = await content_task

        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{template.name}.docx"'
        return response

    except Exception as e:
        await sync_to_async(messages.error)(request, f'Ошибка при генерации документа: {str(e)}')
        return redirect(reverse('document_result', kwargs={'template_id': template_id}))

async def legal_news(request):
    all_news = [
        {
            'title': 'Государственная Дума приняла законопроект о платформенной экономике в первом чтении',
            'content': 'Как новый законопроект защитит права продавцов и покупателей на маркетплейсах...',
            'published_at': '21.07.2025',
            'source': 'Государственная Дума РФ',
            'url': '/news/platform-economy/',
        },
        {
            'title': 'Законопроект о платформенной экономике принят в третьем чтении',
            'content': 'Как новый законопроект защитит права продавцов и покупателей на маркетплейсах...',
            'published_at': '28.07.2025',
            'source': 'Государственная Дума РФ',
            'url': '/news/platform-economy_2/',
        }
    ]
    
    return await sync_to_async(render)(request, 'main/legal_news.html', {'all_news': all_news})

async def privacy_policy(request):
    return await sync_to_async(render)(request, 'main/privacy_policy.html')

async def platform_economy_news_detail(request):
    return await sync_to_async(render)(request, 'main/platform_economy_news_detail.html')

async def platform_economy_news_detail_2(request):
    return await sync_to_async(render)(request, 'main/platform_economy_news_detail_2.html')

async def reference_materials(request):
    query = request.GET.get('q', '').strip()

    if query:
        articles_task = asyncio.create_task(sync_to_async(list)(
            FAQ.objects.filter(
                Q(title__iregex=query) | 
                Q(description__iregex=query) | 
                Q(synonyms__iregex=query),
                is_active=True
            )
        ))
    else:
        articles_task = asyncio.create_task(sync_to_async(list)(FAQ.objects.filter(is_active=True)))

    articles = await articles_task

    for article in articles:
        article.tags_list = [tag.strip() for tag in article.tags.split(',')] if article.tags else []

    context = {
        'articles': articles,
        'query': query,
    }

    return await sync_to_async(render)(request, 'main/reference_materials.html', context)

async def marketplace_rights(request):
    return await sync_to_async(render)(request, 'main/marketplace_rights.html')

async def order_cancellation(request):
    return await sync_to_async(render)(request, 'main/order_cancellation.html')

async def price_error(request):
    return await sync_to_async(render)(request, 'main/price_error.html')

async def glossary(request):
    return await sync_to_async(render)(request, 'main/glossary.html')

async def price_category(request):
    return await sync_to_async(render)(request, 'main/price.html')

async def pretnezya(request):
    return await sync_to_async(render)(request, 'main/pretnezya.html')

async def return_category(request):
    articles_task = asyncio.create_task(sync_to_async(list)(
        FAQ.objects.filter(tags__icontains='Товар') | 
        FAQ.objects.filter(tags__icontains='Качество') | 
        FAQ.objects.filter(tags__icontains='Брак') | 
        FAQ.objects.filter(tags__icontains='Упаковка')
    ))
    articles = await articles_task
    
    for article in articles:
        article.tags_list = [tag.strip() for tag in article.tags.split(',')]
    
    context = {
        'articles': articles,
        'page_title': 'Проблемы с товаром',
        'page_description': 'Решение проблем, связанных с качеством и характеристиками товаров'
    }
    return await sync_to_async(render)(request, 'main/return.html', context)

async def seller_category(request):
    articles_task = asyncio.create_task(sync_to_async(list)(
        FAQ.objects.filter(tags__icontains='Продавец') | 
        FAQ.objects.filter(tags__icontains='Отмена') | 
        FAQ.objects.filter(tags__icontains='Оферта')
    ))
    articles = await articles_task
    
    for article in articles:
        article.tags_list = [tag.strip() for tag in article.tags.split(',')]
    
    context = {
        'articles': articles,
        'page_title': 'Проблемы с продавцом',
        'page_description': 'Решение конфликтов и споров с продавцами на маркетплейсах'
    }
    return await sync_to_async(render)(request, 'main/seller.html', context)

async def delivery_category(request):
    articles_task = asyncio.create_task(sync_to_async(list)(
        FAQ.objects.filter(tags__icontains='Доставка') | 
        FAQ.objects.filter(tags__icontains='Получение')
    ))
    articles = await articles_task
    
    for article in articles:
        article.tags_list = [tag.strip() for tag in article.tags.split(',')]
    
    context = {
        'articles': articles,
        'page_title': 'Проблемы с доставкой',
        'page_description': 'Решение вопросов, связанных с доставкой товаров'
    }
    return await sync_to_async(render)(request, 'main/delivery.html', context)

async def damaged_category(request):
    return await sync_to_async(render)(request, 'main/damaged.html')

async def delivery_terms_category(request):
    return await sync_to_async(render)(request, 'main/delivery_terms.html')

async def bad_goods_delivered(request):
    return await sync_to_async(render)(request, 'main/bad_goods_delivered.html')

async def dosydebnii_isk(request):
    return await sync_to_async(render)(request, 'main/dosydebnii_isk.html')

async def isk_zyavlenie(request):
    return await sync_to_async(render)(request, 'main/isk_zyavlenie.html')

async def return_bad(request):
    return await sync_to_async(render)(request, 'main/return_bad.html')

async def return_good(request):
    return await sync_to_async(render)(request, 'main/return_good.html')

async def return_nevozvrat(request):
    return await sync_to_async(render)(request, 'main/return_nevozvrat.html')

async def return_plata(request):
    return await sync_to_async(render)(request, 'main/return_plata.html')

async def return_vskr_ypakovka(request):
    return await sync_to_async(render)(request, 'main/return_vskr_ypakovka.html')

async def wrong_price(request):
    return await sync_to_async(render)(request, 'main/wrong_price.html')

async def download_empty_template(request, template_id):
    template_task = asyncio.create_task(sync_to_async(get_object_or_404)(DocumentTemplate, id=template_id))
    template = await template_task

    if not template.word_template:
        messages.error(request, "Файл шаблона Word не найден.")
        return redirect('preview_template', template_id=template.id)

    def process_docx():
        doc = Document(template.word_template.path)

        for paragraph in doc.paragraphs:
            for field in template.required_fields:
                placeholder = f'{{{field}}}'
                if placeholder in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if placeholder in inline[i].text:
                            inline[i].text = inline[i].text.replace(placeholder, '        ')

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for field in template.required_fields:
                            placeholder = f'{{{field}}}'
                            if placeholder in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    if placeholder in inline[i].text:
                                        inline[i].text = inline[i].text.replace(placeholder, '        ')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    process_task = asyncio.get_event_loop().run_in_executor(executor, process_docx)
    buffer = await process_task

    filename = f'Пустой_{template.name}.docx'

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

async def download_empty_template_pdf(request, template_id):
    template_task = asyncio.create_task(sync_to_async(get_object_or_404)(DocumentTemplate, id=template_id))
    template = await template_task

    if not template.pdf_template:
        messages.error(request, "PDF шаблон не найден.")
        return redirect('preview_template', template_id=template.id)

    pdf_path = template.pdf_template.path

    if not os.path.exists(pdf_path):
        raise Http404("Файл PDF не найден")

    def read_file():
        with open(pdf_path, 'rb') as f:
            return f.read()

    read_task = asyncio.get_event_loop().run_in_executor(executor, read_file)
    content = await read_task

    response = HttpResponse(content, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Пустой_{template.name}.pdf"'
    return response